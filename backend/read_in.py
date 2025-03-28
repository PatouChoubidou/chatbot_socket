
from pypdf import PdfReader
import os
import docx


async def readAllPDFs(dir='pdfs/'):
    """
    Takes all pdfs in a directory 
    reads in all pdf files into data chunks
    Input:
        - dir: str -- a path to a dir
    Returns:
        - data: array -- an array of data objects full of text and meta data
    """
    data = []
    try:
        arr_pdfs = [x for x in os.listdir(dir) if x.endswith(".pdf")]
    except Exception as e: 
        print(e)
    if(arr_pdfs):
        for pdfUri in arr_pdfs:
            newData = await getDataFromPDF(dir+""+pdfUri)
            data.extend(newData)   
    return data



async def getDataFromPDF(path: os.path):

    """
    Reads in text from a pdf, 
    splits it into paragraphs and smaller chunks
    adds information to it
    and constructs a list of those peaces

    Input:
        - pdf_path: a path to the file
    Returns:
        - a list of data chunks
    """
    
    fileName = os.path.basename(path)
    data = []
    
    with open(path, "rb") as f:
        reader = PdfReader(f)
        allpages = reader.pages
        
        # read in text 
        for pageNum, page in enumerate(allpages):
            pageTxt = page.extract_text()

            # split each page into paragraph chunks for a more accurate rag
            # the size of the splitted chunks has effect on the output
           
            # paragraphs = re.split('\s{4,}', pageTxt)
            # paragraphs = pageTxt.split("\n\n")
            # paragraphs = pageTxt.split("\n ")
            
            # in this case the paragraph is more like a sentence aka small chunks :-)
            paragraphs = pageTxt.split(". ")
            for paraIndex, paragraph in enumerate(paragraphs):
                
                newTextObj = {}
                newTextObj["filename"] = fileName
                newTextObj["content"] = paragraph
                newTextObj["page"] = pageNum
                newTextObj["paragraph"] = paraIndex + 1
                    
                data.append(newTextObj)   
            
    return data


# docx not in use yet
async def getTextFromDocx(path: os.path):
    """
    Opens a .docx or doc file and returns its txt
    Params:
        path: os.path - filepath
    Returns:
        - a list of text data chunks
    """
    filename = os.path.basename(path)
    data = []
    try:
        doc = docx.Document(path)
       
        for paraIndex, para in enumerate(doc.paragraphs):

            print(f"paragaph from file {para}")
           
            newTextObj = {}
            newTextObj["filename"] = filename
            newTextObj["content"] = para.text
            newTextObj["page"] = 0
            newTextObj["paragraph"] = paraIndex + 1
                    
            data.append(newTextObj)   
        
    except:
        print("Could read in the document")

    return data
    

async def getTextFromTxt(path: os.path):
    """
    Opens a .docx or doc file and returns its txt
    Params:
        path: os.path - filepath
    Returns:
        - a list of text data chunks
    """
    filename = os.path.basename(path)
    data = []

    with open(path) as file_in:
        lines = []
        for line in file_in:
            lines.append(line)
            
    for lineIndex, line in enumerate(lines):
           
            newTextObj = {}
            newTextObj["filename"] = filename
            newTextObj["content"] = line
            newTextObj["page"] = 0
            newTextObj["paragraph"] = lineIndex + 1
                    
            data.append(newTextObj)   

    return data


async def getFullTextFromPDF(path: os.path):

    """
    Reads in text from a pdf

    Input:
        - pdf_path: a path to the file
    Returns:
        - a list of data chunks
    """
    
    fileName = os.path.basename(path)
    out_txt = ""
    
    with open(path, "rb") as f:
        reader = PdfReader(f)
        allpages = reader.pages
        
        # read in text 
        for pageNum, page in enumerate(allpages):
            pageTxt = page.extract_text()
            out_txt += pageTxt       
            
    return out_txt



async def getFullTextFromDoc(path: os.path):
    """
    Opens a .docx or doc file and returns its txt
    Params:
        path: os.path - filepath
    Returns:
        - the text of a document
    """
    fullText = []
    try:
        doc = docx.Document(path)
        print(doc)
       
        for para in doc.paragraphs:
            fullText.append(para.text)
    
    except:
        print("Could read in the document")

    return '\n'.join(fullText)


async def getFullTextFromTxt(path: os.path):
    """
    Opens a .docx or doc file and returns its txt
    Params:
        path: os.path - filepath
    Returns:
        - a list of text data chunks
    """
    filename = os.path.basename(path)
    lines = []
    
    with open(path) as file_in:
        for line in file_in:
            lines.append(line)
            
    return '\n'.join(lines)





